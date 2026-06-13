from docx import Document

# Create document
document = Document()

# Add heading
document.add_heading(
    'Python-Docx Example',
    level=1
)

# Add paragraph
document.add_paragraph(
    'This Word document was created using Python.'
)

# Add another paragraph
document.add_paragraph(
    'Python-Docx makes document automation easy.'
)

# Save document
document.save("sample.docx")

print("Word document created successfully! ✅")